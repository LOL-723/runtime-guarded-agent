import hashlib
import json
import logging
import math
import re
import threading
from pathlib import Path

from fastapi import UploadFile

from core.config import settings
from llm.client import llm_client
from schemas.rag import RagAskResponse, RagChunk, RagSource, RagUploadResponse


SUPPORTED_RAG_SUFFIXES = {".txt", ".md", ".csv", ".json", ".pdf", ".docx"}


class RagError(ValueError):
    pass


class RagService:
    def __init__(self) -> None:
        self.storage_dir = Path(settings.RAG_STORAGE_DIR)
        self.max_upload_bytes = settings.RAG_MAX_UPLOAD_MB * 1024 * 1024
        self.chunk_size = settings.RAG_CHUNK_SIZE
        self.chunk_overlap = min(settings.RAG_CHUNK_OVERLAP, max(0, self.chunk_size - 1))
        self._lock = threading.Lock()
        self._embedding_model = None
        self._rerank_model = None
        self._collection = None

    async def upload(self, file: UploadFile) -> RagUploadResponse:
        filename = Path(file.filename or "").name
        if not filename:
            raise RagError("filename cannot be empty")

        suffix = Path(filename).suffix.lower()
        if suffix not in SUPPORTED_RAG_SUFFIXES:
            allowed = ", ".join(sorted(SUPPORTED_RAG_SUFFIXES))
            raise RagError(f"unsupported file type: {suffix}. allowed: {allowed}")

        content = await file.read()
        if not content:
            raise RagError("uploaded file cannot be empty")
        if len(content) > self.max_upload_bytes:
            raise RagError(f"file too large. max upload size is {settings.RAG_MAX_UPLOAD_MB}MB")

        document_id = hashlib.sha256(content).hexdigest()[:16]
        doc_dir = self.storage_dir / document_id
        doc_dir.mkdir(parents=True, exist_ok=True)
        file_path = doc_dir / self._safe_filename(filename)
        file_path.write_bytes(content)

        text = self.extract_text(file_path)
        if not text.strip():
            raise RagError("no text could be extracted from this file")

        chunk_texts = self.split_into_chunks(text)
        if not chunk_texts:
            raise RagError("no chunks could be created from this file")

        embeddings = self.embed_chunks(chunk_texts)
        chunk_ids = [f"{document_id}-{index:04d}" for index in range(len(chunk_texts))]
        metadatas = [
            {"document_id": document_id, "filename": filename, "chunk_index": index}
            for index in range(len(chunk_texts))
        ]

        try:
            self.collection.delete(where={"document_id": document_id})
        except Exception:
            pass

        self.collection.upsert(
            ids=chunk_ids,
            embeddings=embeddings,
            metadatas=metadatas,
            documents=chunk_texts,
        )

        chunks = [
            RagChunk(id=chunk_ids[index], index=index, content=chunk)
            for index, chunk in enumerate(chunk_texts)
        ]
        (doc_dir / "chunks.json").write_text(
            json.dumps([chunk.model_dump() for chunk in chunks], ensure_ascii=False, indent=2),
            encoding="utf-8",
        )

        return RagUploadResponse(
            document_id=document_id,
            filename=filename,
            chunk_count=len(chunks),
            chunks=chunks,
        )

    def ask(self, question: str, top_k: int | None = None, rerank_top_k: int | None = None) -> RagAskResponse:
        if not question or not question.strip():
            raise RagError("question cannot be empty")

        retrieved = self.retrieve(question, top_k or settings.RAG_RETRIEVE_TOP_K)
        if not retrieved:
            return RagAskResponse(answer="根据已上传文档无法回答。", sources=[])

        reranked = self.rerank(question, retrieved, rerank_top_k or settings.RAG_RERANK_TOP_K)
        answer = self.generate(question, [source.content for source in reranked])
        return RagAskResponse(answer=answer, sources=reranked)

    def retrieve(self, query: str, top_k: int) -> list[RagSource]:
        query_embedding = self.embed_text(query)
        results = self.collection.query(query_embeddings=[query_embedding], n_results=top_k)
        documents = (results.get("documents") or [[]])[0]
        metadatas = (results.get("metadatas") or [[]])[0]
        ids = (results.get("ids") or [[]])[0]

        sources = []
        for chunk_id, content, metadata in zip(ids, documents, metadatas):
            metadata = metadata or {}
            sources.append(
                RagSource(
                    chunk_id=chunk_id,
                    document_id=str(metadata.get("document_id", "")),
                    filename=str(metadata.get("filename", "")),
                    chunk_index=int(metadata.get("chunk_index", 0)),
                    content=content,
                )
            )
        return sources

    def rerank(self, query: str, sources: list[RagSource], top_k: int) -> list[RagSource]:
        if not sources:
            return []
        if self.rerank_model is None:
            return sources[:top_k]
        pairs = [(query, source.content) for source in sources]
        scores = self.rerank_model.predict(pairs)
        ranked = sorted(zip(sources, scores), key=lambda item: float(item[1]), reverse=True)
        return [source for source, _ in ranked[:top_k]]

    def generate(self, question: str, chunks: list[str]) -> str:
        context = "\n\n".join(f"[{index}] {chunk}" for index, chunk in enumerate(chunks))
        system_prompt = (
            "你是一个知识库问答助手。只能根据提供的文档片段回答。"
            "如果片段中没有答案，直接说：根据已上传文档无法回答。"
        )
        user_message = f"""用户问题:
{question}

相关片段:
{context}

请基于相关片段作答，不要编造信息。"""
        return llm_client.chat(user_message=user_message, system_prompt=system_prompt)

    def split_into_chunks(self, text: str) -> list[str]:
        normalized = re.sub(r"\r\n?", "\n", text).strip()
        normalized = re.sub(r"\n{3,}", "\n\n", normalized)
        chunks = []
        start = 0
        while start < len(normalized):
            hard_end = min(start + self.chunk_size, len(normalized))
            end = self._find_chunk_end(normalized, start, hard_end)
            chunk = normalized[start:end].strip()
            if chunk:
                chunks.append(chunk)
            if end >= len(normalized):
                break
            start = max(end - self.chunk_overlap, start + 1)
        return chunks

    def embed_text(self, text: str) -> list[float]:
        if self.embedding_model is None:
            return self._hash_embedding(text)
        embedding = self.embedding_model.encode(text, normalize_embeddings=True)
        return embedding.tolist()

    def embed_chunks(self, chunks: list[str]) -> list[list[float]]:
        if self.embedding_model is None:
            return [self._hash_embedding(chunk) for chunk in chunks]
        embeddings = self.embedding_model.encode(chunks, normalize_embeddings=True)
        return embeddings.tolist()

    def extract_text(self, file_path: Path) -> str:
        suffix = file_path.suffix.lower()
        if suffix in {".txt", ".md", ".csv", ".json"}:
            return self._read_text(file_path)
        if suffix == ".pdf":
            return self._extract_pdf(file_path)
        if suffix == ".docx":
            return self._extract_docx(file_path)
        raise RagError(f"unsupported file type: {suffix}")

    @property
    def embedding_model(self):
        if self._embedding_model is None:
            with self._lock:
                if self._embedding_model is None:
                    try:
                        from sentence_transformers import SentenceTransformer

                        self._embedding_model = SentenceTransformer(settings.RAG_EMBEDDING_MODEL)
                    except Exception as exc:
                        logging.warning(
                            "Failed to load embedding model %s; using local hash embeddings: %s",
                            settings.RAG_EMBEDDING_MODEL,
                            exc,
                        )
                        self._embedding_model = False
        if self._embedding_model is False:
            return None
        return self._embedding_model

    @property
    def rerank_model(self):
        if self._rerank_model is None:
            with self._lock:
                if self._rerank_model is None:
                    try:
                        from sentence_transformers import CrossEncoder

                        self._rerank_model = CrossEncoder(settings.RAG_RERANK_MODEL)
                    except Exception as exc:
                        logging.warning(
                            "Failed to load rerank model %s; skipping rerank: %s",
                            settings.RAG_RERANK_MODEL,
                            exc,
                        )
                        self._rerank_model = False
        if self._rerank_model is False:
            return None
        return self._rerank_model

    @property
    def collection(self):
        if self._collection is None:
            with self._lock:
                if self._collection is None:
                    import chromadb

                    Path(settings.RAG_CHROMA_DIR).mkdir(parents=True, exist_ok=True)
                    client = chromadb.PersistentClient(path=settings.RAG_CHROMA_DIR)
                    self._collection = client.get_or_create_collection(name=settings.RAG_COLLECTION_NAME)
        return self._collection

    def _find_chunk_end(self, text: str, start: int, hard_end: int) -> int:
        if hard_end >= len(text):
            return len(text)
        window = text[start:hard_end]
        for delimiter in ("\n\n", "\n", "。", "！", "？", ".", "!", "?"):
            pos = window.rfind(delimiter)
            if pos >= self.chunk_size * 0.5:
                return start + pos + len(delimiter)
        return hard_end

    @staticmethod
    def _read_text(file_path: Path) -> str:
        for encoding in ("utf-8", "utf-8-sig", "gb18030"):
            try:
                return file_path.read_text(encoding=encoding)
            except UnicodeDecodeError:
                continue
        return file_path.read_text(encoding="utf-8", errors="ignore")

    @staticmethod
    def _extract_pdf(file_path: Path) -> str:
        from pypdf import PdfReader

        reader = PdfReader(str(file_path))
        pages = []
        for page_index, page in enumerate(reader.pages, start=1):
            page_text = page.extract_text() or ""
            if page_text.strip():
                pages.append(f"[page {page_index}]\n{page_text}")
        return "\n\n".join(pages)

    @staticmethod
    def _extract_docx(file_path: Path) -> str:
        from docx import Document

        doc = Document(str(file_path))
        paragraphs = [paragraph.text for paragraph in doc.paragraphs if paragraph.text.strip()]
        table_rows = []
        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                if any(cells):
                    table_rows.append(" | ".join(cells))
        return "\n\n".join(paragraphs + table_rows)

    @staticmethod
    def _safe_filename(filename: str) -> str:
        stem = Path(filename).stem
        suffix = Path(filename).suffix.lower()
        safe_stem = re.sub(r"[^A-Za-z0-9._-]+", "_", stem).strip("._")
        return f"{safe_stem or 'document'}{suffix}"

    @staticmethod
    def _hash_embedding(text: str, dimensions: int = 384) -> list[float]:
        vector = [0.0] * dimensions
        tokens = re.findall(r"[\w\u4e00-\u9fff]+", text.lower())
        if not tokens:
            tokens = [text]
        for token in tokens:
            digest = hashlib.sha256(token.encode("utf-8")).digest()
            for offset in range(0, len(digest), 2):
                index = int.from_bytes(digest[offset:offset + 2], "little") % dimensions
                sign = 1.0 if digest[offset] % 2 == 0 else -1.0
                vector[index] += sign
        norm = math.sqrt(sum(value * value for value in vector)) or 1.0
        return [value / norm for value in vector]


rag_service = RagService()
